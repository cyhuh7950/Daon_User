from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "superpowers" / "specs" / "2026-07-20-daon-user-program-design.md"
OUTPUT = ROOT / "docs" / "Daon 사용자형 지식 업무지원 프로그램 상세 설계서.docx"
ARCH_IMAGE = ROOT / "docs" / ".build" / "daon_user_architecture.png"


# compact_reference_guide preset + named Korean typography/cover overrides
FONT_BODY = "Malgun Gothic"
FONT_MONO = "Consolas"
COLOR_NAVY = "203748"
COLOR_BLUE = "2E74B5"
COLOR_DARK_BLUE = "1F4D78"
COLOR_MUTED = "5B6573"
COLOR_LIGHT = "F4F6F9"
COLOR_TABLE_HEADER = "E8EEF5"
COLOR_TABLE_BORDER = "C8D2DE"
COLOR_GOLD = "9A6A00"
COLOR_WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_SIDE = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, *, name: str = FONT_BODY, size: float | None = None,
                 color: str | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name: str, size: float, color: str | None = None,
                   bold: bool | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    if color:
        style.font.color.rgb = rgb(color)
    if bold is not None:
        style.font.bold = bold


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = CELL_MARGIN_TOP_BOTTOM,
                     start: int = CELL_MARGIN_SIDE,
                     bottom: int = CELL_MARGIN_TOP_BOTTOM,
                     end: int = CELL_MARGIN_SIDE) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = COLOR_TABLE_BORDER, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: Sequence[int], *, indent: int = TABLE_INDENT_DXA,
                       border_color: str = COLOR_TABLE_BORDER) -> None:
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"Table widths must total {CONTENT_DXA}, got {sum(widths)}")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_table_borders(table, border_color)


def table_widths(headers: Sequence[str], rows: Sequence[Sequence[str]]) -> list[int]:
    cols = len(headers)
    if cols == 1:
        return [CONTENT_DXA]
    if cols == 2:
        left_max = max([len(headers[0])] + [len(row[0]) for row in rows if row])
        return [2160, 7200] if left_max <= 22 else [3240, 6120]
    if cols == 3:
        return [1440, 3600, 4320]
    if cols == 4:
        return [1260, 2520, 2880, 2700]
    base = CONTENT_DXA // cols
    widths = [base] * cols
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=COLOR_MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, value, fld_end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    set_style_font(normal, FONT_BODY, 11, "20242A")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, COLOR_BLUE, 18, 10),
        "Heading 2": (13, COLOR_BLUE, 14, 7),
        "Heading 3": (12, COLOR_DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        set_style_font(style, FONT_BODY, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run("Daon 사용자형 지식 업무지원 프로그램")
    set_run_font(left, size=8.5, color=COLOR_MUTED, bold=True)
    right = p.add_run("\t상세 설계 정본")
    set_run_font(right, size=8.5, color=COLOR_MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), COLOR_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), FONT_BODY)
    rpr.extend([rfonts, color, underline])
    run.append(rpr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_TOKEN = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\)|<https?://[^>]+>)")


def add_inline_runs(paragraph, text: str, *, size: float = 11, color: str = "20242A") -> None:
    position = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name=FONT_MONO, size=max(size - 0.5, 8), color=COLOR_DARK_BLUE)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "EEF2F6")
            run._element.get_or_add_rPr().append(shd)
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            url = token[1:-1]
            add_hyperlink(paragraph, url, url)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color)


def next_abstract_num_id(numbering) -> int:
    values = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    return max(values, default=-1) + 1


def next_num_id(numbering) -> int:
    values = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    return max(values, default=0) + 1


def add_abstract_numbering(doc: Document, *, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = next_abstract_num_id(numbering)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    bullet_marks = ["•", "◦", "▪"]
    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), f"%{level + 1}." if ordered else bullet_marks[level])
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        left = 540 + level * 540
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), "270")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.extend([tabs, ind, spacing])
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        for attr in ("ascii", "hAnsi", "eastAsia"):
            rfonts.set(qn(f"w:{attr}"), FONT_BODY)
        rpr.append(rfonts)
        lvl.extend([start, num_fmt, lvl_text, suff, ppr, rpr])
        abstract.append(lvl)
    numbering.insert(0, abstract)
    return abstract_id


def create_num_instance(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    num_id = next_num_id(numbering)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 2)))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_callout(doc: Document, title: str, body_lines: Iterable[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA], border_color="D6DFE8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLOR_LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_run_font(run, size=10.5, color=COLOR_DARK_BLUE, bold=True)
    for line in body_lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.08)
        add_inline_runs(p, line, size=9.7)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(110)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("독립 제품 상세 설계")
    set_run_font(r, size=11, color=COLOR_GOLD, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Daon 사용자형 지식 업무지원 프로그램")
    set_run_font(r, size=28, color=COLOR_NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(44)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Source Workspace · 선택형 LLM · 업무 Studio")
    set_run_font(r, size=14, color=COLOR_DARK_BLUE)

    for label, value in (
        ("문서 구분", "독립 제품 상세 설계 정본"),
        ("작성일", "2026-07-20"),
        ("상태", "승인 · 신산님 · 2026-07-20"),
        ("제품 관계", "Daon2 · Daon2.5 · Daon3과 별개의 독립 제품"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{label}  ")
        set_run_font(r, size=10, color=COLOR_MUTED, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10, color=COLOR_MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(45)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("설계 정본은 Markdown이며 이 DOCX는 승인·배포용 생성본입니다.")
    set_run_font(r, size=9, color=COLOR_MUTED, italic=True)
    p.add_run().add_break(WD_BREAK.PAGE)

    add_callout(
        doc,
        "설계 핵심",
        [
            "Daon 계열과 별개의 자체 UI·API·DB·배포·운영 제품",
            "문서·표·이미지는 Vision/LLM-first, 오디오는 Audio LLM 또는 ASR+LLM 의미 이해, 보조 추출만으로 완료 금지",
            "Daon 승인 지식과 강제 RuleSet 우선, 사용자 가중치와 충돌 공개",
            "로컬 LLM·사내 LLM·외부 LLM을 정책에 따라 선택",
            "생성 설정을 확인한 뒤 업무 산출물을 생성·검토·승인·전달",
            "로컬 비공개와 클라우드 동기화, Web·Windows·iOS·Android 지원",
        ],
    )


def load_font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/malgunbd.ttf" if bold else "C:/Windows/Fonts/malgun.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str,
                       font, fill: str = "#203748") -> None:
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=5, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(((x1 + x2 - width) / 2, (y1 + y2 - height) / 2), text,
                        font=font, fill=fill, spacing=5, align="center")


def make_architecture_image() -> None:
    image = Image.new("RGB", (1600, 1020), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(42, True)
    box_font = load_font(27, True)
    small_font = load_font(22, False)
    draw.text((80, 45), "Daon 사용자 프로그램 논리 아키텍처", font=title_font, fill="#203748")

    boxes: dict[str, tuple[int, int, int, int]] = {}

    def box(key: str, xy, text: str, fill: str, font=box_font):
        boxes[key] = xy
        draw.rounded_rectangle(xy, radius=18, fill=fill, outline="#8CA0B5", width=3)
        draw_centered_text(draw, xy, text, font)

    def arrow(a: str, b: str):
        ax1, ay1, ax2, ay2 = boxes[a]
        bx1, by1, bx2, by2 = boxes[b]
        start = ((ax1 + ax2) // 2, ay2)
        end = ((bx1 + bx2) // 2, by1)
        draw.line([start, end], fill="#60758A", width=5)
        draw.polygon([(end[0], end[1]), (end[0] - 10, end[1] - 18), (end[0] + 10, end[1] - 18)], fill="#60758A")

    box("cloud_clients", (120, 130, 950, 230), "Web · iOS · Android · Windows Cloud", "#E8EEF5", small_font)
    box("windows_local", (1000, 130, 1480, 230), "Windows Tauri\nLocal-private", "#E8F4F1", small_font)
    box("gateway", (380, 290, 1220, 390), "Web BFF · Public API Gateway", "#DDEAF6")
    arrow("cloud_clients", "gateway")
    arrow("windows_local", "gateway")

    core_items = [
        ("iam", (80, 470, 360, 590), "Identity\nTenant · Policy"),
        ("workspace", (390, 470, 690, 590), "Workspace\nSource"),
        ("knowledge", (720, 470, 1020, 590), "Knowledge\nRetrieval"),
        ("run", (1050, 470, 1360, 590), "Run · Studio\nReview"),
    ]
    for key, xy, text in core_items:
        box(key, xy, text, "#F4F6F9", small_font)
        arrow("gateway", key)

    box("data", (80, 700, 540, 820), "Cloud Data\nPostgreSQL · Vector · Object", "#F0F5E8", small_font)
    box("model", (575, 700, 1025, 820), "Model Gateway\nLocal · Internal · External", "#FFF3DA", small_font)
    box("connector", (1060, 700, 1520, 820), "Connector Layer\nDaon API · Internet", "#FDEBEC", small_font)
    arrow("workspace", "data")
    arrow("knowledge", "model")
    arrow("run", "connector")

    box("local", (300, 890, 1300, 980), "Local Workspace Engine · Encrypted SQLite · Files · Vector Index · Local LLM", "#E8F4F1", small_font)
    # Gateway sync path and the Windows-only Loopback Local API path are distinct.
    draw.line([(800, 390), (60, 430), (60, 930), (300, 930)], fill="#60758A", width=4)
    draw.polygon([(300, 930), (282, 920), (282, 940)], fill="#60758A")
    draw.text((75, 845), "Approved Sync", font=small_font, fill="#60758A")
    draw.line([(1240, 230), (1540, 270), (1540, 930), (1300, 930)], fill="#60758A", width=4)
    draw.polygon([(1300, 930), (1318, 920), (1318, 940)], fill="#60758A")
    draw.text((1260, 845), "Loopback\nLocal API", font=small_font, fill="#60758A")

    ARCH_IMAGE.parent.mkdir(parents=True, exist_ok=True)
    image.save(ARCH_IMAGE, dpi=(180, 180))


def set_picture_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_architecture_figure(doc: Document) -> None:
    make_architecture_image()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    shape = run.add_picture(str(ARCH_IMAGE), width=Inches(6.35))
    set_picture_alt_text(
        shape,
        "Daon 사용자 프로그램 논리 아키텍처",
        "Web·모바일·Windows Cloud 경로는 BFF와 공개 API를 사용하고 Windows Local-private는 Loopback Local API로 Local Workspace Engine을 직접 사용하는 구조",
    )
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    r = caption.add_run("그림 1. 독립 서비스·모델·Connector·로컬 실행 경계")
    set_run_font(r, size=8.8, color=COLOR_MUTED, italic=True)


def add_code_block(doc: Document, lines: Sequence[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA], border_color="D6DFE8")
    set_table_row_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F8FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name=FONT_MONO, size=8.5, color="263442")
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def clean_table_cell(value: str) -> str:
    return value.strip().replace("`", "")


def add_markdown_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, COLOR_TABLE_HEADER)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True
        add_inline_runs(p, clean_table_cell(header), size=9, color=COLOR_NAVY)
        for run in p.runs:
            run.bold = True
    set_repeat_table_header(table.rows[0])
    set_table_row_cant_split(table.rows[0])

    for row_values in rows:
        row = table.add_row()
        set_table_row_cant_split(row)
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if len(clean_table_cell(value)) <= 16 and idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, clean_table_cell(value), size=8.8)

    set_table_geometry(table, table_widths(headers, rows))
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)


def is_table_separator(line: str) -> bool:
    parts = [part.strip() for part in line.strip().strip("|").split("|")]
    return bool(parts) and all(re.fullmatch(r":?-{3,}:?", part) for part in parts)


def parse_table_line(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def add_heading(doc: Document, level: int, text: str) -> None:
    word_level = min(max(level - 1, 1), 3)  # Markdown ## -> Word Heading 1
    p = doc.add_paragraph(style=f"Heading {word_level}")
    add_inline_runs(p, text, size={1: 16, 2: 13, 3: 12}[word_level],
                    color={1: COLOR_BLUE, 2: COLOR_BLUE, 3: COLOR_DARK_BLUE}[word_level])
    for run in p.runs:
        run.bold = True


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    add_inline_runs(p, text)


def build_from_markdown(doc: Document, markdown_text: str) -> None:
    lines = markdown_text.splitlines()
    start = next(i for i, line in enumerate(lines) if line.strip() == "## 0. 설계 기준")
    lines = lines[start:]

    bullet_abstract = add_abstract_numbering(doc, ordered=False)
    decimal_abstract = add_abstract_numbering(doc, ordered=True)
    bullet_num = create_num_instance(doc, bullet_abstract)
    active_decimal_num: int | None = None
    in_code = False
    code_language = ""
    code_lines: list[str] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            add_body_paragraph(doc, " ".join(line.strip() for line in paragraph_lines))
            paragraph_lines = []

    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if in_code:
            if stripped.startswith("```"):
                if code_language == "mermaid":
                    add_architecture_figure(doc)
                else:
                    add_code_block(doc, code_lines)
                in_code = False
                code_language = ""
                code_lines = []
            else:
                code_lines.append(raw)
            i += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            in_code = True
            code_language = stripped[3:].strip().lower()
            i += 1
            continue

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if stripped == "---":
            flush_paragraph()
            doc.add_page_break()
            active_decimal_num = None
            i += 1
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            add_heading(doc, len(heading.group(1)), heading.group(2).strip())
            active_decimal_num = None
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            flush_paragraph()
            headers = parse_table_line(stripped)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = parse_table_line(lines[i])
                if len(row) < len(headers):
                    row.extend([""] * (len(headers) - len(row)))
                rows.append(row[:len(headers)])
                i += 1
            add_markdown_table(doc, headers, rows)
            active_decimal_num = None
            continue

        unordered = re.match(r"^(\s*)-\s+(.+)$", raw)
        if unordered:
            flush_paragraph()
            level = min(len(unordered.group(1).replace("\t", "    ")) // 2, 2)
            p = doc.add_paragraph()
            apply_numbering(p, bullet_num, level)
            add_inline_runs(p, unordered.group(2).strip())
            active_decimal_num = None
            i += 1
            continue

        ordered = re.match(r"^(\s*)\d+\.\s+(.+)$", raw)
        if ordered:
            flush_paragraph()
            if active_decimal_num is None:
                active_decimal_num = create_num_instance(doc, decimal_abstract)
            level = min(len(ordered.group(1).replace("\t", "    ")) // 2, 2)
            p = doc.add_paragraph()
            apply_numbering(p, active_decimal_num, level)
            add_inline_runs(p, ordered.group(2).strip())
            i += 1
            continue

        paragraph_lines.append(raw)
        i += 1

    flush_paragraph()


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Daon 사용자형 지식 업무지원 프로그램 상세 설계서"
    props.subject = "독립 제품 상세 설계 정본"
    props.author = "Daon 설계 책임자"
    props.keywords = "Daon, 지식 업무지원, 로컬 LLM, RAG, 업무 Studio"
    props.comments = "2026-07-20 TP-0 승인과 운영·화면·same-origin API·3단계 배포·Subagent 의무 보고 계약을 반영한 배포용 생성본"


def main() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    set_core_properties(doc)
    add_cover(doc)
    build_from_markdown(doc, text)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
