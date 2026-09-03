from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DIST = ROOT / "docs" / "manual" / "dist"
PUBLIC = ROOT / "apps" / "web" / "public" / "manual"
RELEASE = "1.0.0"
RELEASED_AT = "2026-08-20"
DOCS = (
    ("daon-getting-started", "Daon Getting Started", "Notebook의 첫 작업 흐름과 안전한 운영"),
    ("daon-user-manual", "Daon 사용자 설명서", "Workspace 화면·설정·운영 절차"),
    ("daon-knowledge-llm-guide", "Daon 지식·LLM 활용 가이드", "지식·Provider·Citation 품질 운영 기준"),
)
MIMES = {
    "markdown": "text/markdown; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}
EXTENSIONS = {"markdown": "md", "docx": "docx", "pdf": "pdf"}
LATIN_FONT = "Noto Sans"
EAST_ASIA_FONT = "Gulim"


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=11, bold=None, color=None):
    run.font.name = LATIN_FONT
    run_fonts = run._element.get_or_add_rPr().rFonts
    run_fonts.set(qn("w:ascii"), LATIN_FONT)
    run_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    run_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_bookmark(paragraph, name, bookmark_id):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_link(paragraph, label, anchor):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "16")
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), "16")
    rpr.extend((fonts, color, underline, size, size_cs))
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.add_run("DAON · ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "222222", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = document.styles[name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        # LibreOffice's synthetic CJK bold can fill complex Hangul counters and
        # produce unreadable black blocks in exported PDFs. Size and colour
        # provide the hierarchy while keeping Korean glyphs legible.
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    if "Daon Caption" not in document.styles:
        style = document.styles.add_style("Daon Caption", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor(90, 90, 100)
        style.paragraph_format.space_after = Pt(8)


def configure_section(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.text = "DAON · USER GUIDE"
    set_run_font(header.runs[0], 9, bold=True, color=(46, 116, 181))
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer)
    for run in footer.runs:
        set_run_font(run, 9, color=(90, 90, 100))


def add_metadata_table(document, title):
    table = document.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.85)
    table.columns[1].width = Inches(4.65)
    table_width = table._tbl.tblPr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table._tbl.tblPr.append(table_width)
    table_width.set(qn("w:w"), "9360")
    table_width.set(qn("w:type"), "dxa")
    values = (("문서", title), ("Release", RELEASE), ("언어", "한국어 (ko-KR)"), ("범위", "공개 안내 + 로그인 후 조직 전용 절차 구분"))
    for row, (label, value) in zip(table.rows, values):
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[0].width = Inches(1.85)
        row.cells[1].width = Inches(4.65)
        row.cells[0].text = label
        row.cells[1].text = value
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "E8EEF5")
        row.cells[0]._tc.get_or_add_tcPr().append(shade)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, 9, bold=True, color=(31, 77, 120))
        for run in row.cells[1].paragraphs[0].runs:
            set_run_font(run, 9)
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(header_flag)


def parse_headings(lines):
    headings = []
    for line in lines:
        match = re.match(r"^(#{2,3})\s+(.+)$", line)
        if match:
            headings.append((len(match.group(1)), match.group(2).strip(), f"toc_{len(headings) + 1}"))
    return headings


def add_cover(document, title, summary):
    document.add_paragraph("DAON RELEASE 1", style=None).alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker = document.paragraphs[-1].runs[0]
    set_run_font(kicker, 10, bold=True, color=(46, 116, 181))
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(72)
    set_run_font(paragraph.add_run(title), 28, bold=True, color=(31, 38, 52))
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(12)
    set_run_font(subtitle.add_run(summary), 12, color=(90, 90, 100))
    line = document.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_before = Pt(48)
    set_run_font(line.add_run(f"Release {RELEASE} · {RELEASED_AT} · 한국어"), 10, bold=True, color=(46, 116, 181))
    document.add_page_break()


def add_toc(document, headings):
    document.add_heading("문서 정보", level=1)
    add_metadata_table(document, document.core_properties.title)
    document.add_heading("목차", level=1)
    for level, title, anchor in headings:
        if level != 2:
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 0.9
        add_internal_link(paragraph, title, anchor)
    document.add_page_break()


def new_numbering(document, ordered):
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if ordered else "•")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "space")
    paragraph_properties = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:start"), "540")
    indent.set(qn("w:hanging"), "271")
    paragraph_properties.append(indent)
    level.append(start)
    level.append(num_fmt)
    level.append(level_text)
    level.append(suffix)
    level.append(paragraph_properties)
    abstract.append(level)
    numbering.append(abstract)
    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    reference = OxmlElement("w:abstractNumId")
    reference.set(qn("w:val"), str(abstract_id))
    number.append(reference)
    numbering.append(number)
    return num_id


def clean_inline(text):
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text.replace("`", ""))


def add_list_paragraph(document, text, num_id):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    reference = OxmlElement("w:numId")
    reference.set(qn("w:val"), str(num_id))
    num_pr.extend((level, reference))
    paragraph._p.get_or_add_pPr().append(num_pr)
    paragraph.add_run(clean_inline(text))
    return paragraph


def add_markdown_body(document, source, lines, headings):
    heading_cursor = 0
    list_kind = None
    list_num_id = None
    index = 0
    while index < len(lines):
        raw = lines[index].rstrip()
        if not raw or raw.startswith("# "):
            index += 1
            continue
        heading = re.match(r"^(#{2,3})\s+(.+)$", raw)
        if heading:
            list_kind = None
            level = len(heading.group(1)) - 1
            paragraph = document.add_heading(heading.group(2).strip(), level=level)
            add_bookmark(paragraph, headings[heading_cursor][2], heading_cursor + 1)
            heading_cursor += 1
            index += 1
            continue
        image = re.match(r"^!\[(.+)]\((.+)\)$", raw)
        if image:
            list_kind = None
            image_path = (source.parent / image.group(2)).resolve()
            if image_path.is_file() and ROOT in image_path.parents:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                inline = run.add_picture(str(image_path), width=Inches(6.35))
                doc_pr = inline._inline.docPr
                doc_pr.set("descr", image.group(1))
                caption = document.add_paragraph(image.group(1), style="Daon Caption")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            index += 1
            continue
        quote = re.match(r"^>\s*(.+)$", raw)
        if quote:
            list_kind = None
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(clean_inline(quote.group(1)))
            set_run_font(run, 10, color=(31, 77, 120))
            paragraph._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))
            index += 1
            continue
        ordered = re.match(r"^\d+\.\s+(.+)$", raw)
        bullet = re.match(r"^-\s+(.+)$", raw)
        if ordered or bullet:
            current_kind = "ordered" if ordered else "bullet"
            if list_kind != current_kind:
                list_kind = current_kind
                list_num_id = new_numbering(document, bool(ordered))
            add_list_paragraph(document, (ordered or bullet).group(1), list_num_id)
            index += 1
            continue
        list_kind = None
        paragraph = document.add_paragraph(clean_inline(raw))
        paragraph.paragraph_format.widow_control = True
        index += 1


def build_docx():
    DIST.mkdir(parents=True, exist_ok=True)
    for document_id, title, summary in DOCS:
        source = ROOT / "docs" / "manual" / document_id / "index.md"
        lines = source.read_text(encoding="utf-8").splitlines()
        headings = parse_headings(lines)
        document = Document()
        document.core_properties.title = title
        document.core_properties.subject = summary
        document.core_properties.author = "Daon"
        document.core_properties.keywords = "Daon, 사용자 설명서, Release 1"
        document.core_properties.comments = "공개 안내와 로그인 후 조직 전용 절차를 구분한 한국어 정본"
        configure_styles(document)
        configure_section(document)
        add_cover(document, title, summary)
        add_toc(document, headings)
        add_markdown_body(document, source, lines, headings)
        document.save(DIST / f"{document_id}.docx")


def sha256(path):
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def build_manifest():
    PUBLIC.mkdir(parents=True, exist_ok=True)
    documents = []
    for document_id, title, summary in DOCS:
        source_md = ROOT / "docs" / "manual" / document_id / "index.md"
        public_md = PUBLIC / f"{document_id}.md"
        shutil.copyfile(source_md, public_md)
        for extension in ("docx", "pdf"):
            shutil.copyfile(DIST / f"{document_id}.{extension}", PUBLIC / f"{document_id}.{extension}")
        assets = {}
        for kind, extension in EXTENSIONS.items():
            path = PUBLIC / f"{document_id}.{extension}"
            assets[kind] = {
                "filename": path.name,
                "href": f"/manual/{path.name}",
                "bytes": path.stat().st_size,
                "sha256": sha256(path),
                "mime": MIMES[kind],
            }
        documents.append({
            "document_id": document_id,
            "title": title,
            "summary": summary,
            "auth_scope": "public_and_authenticated",
            "version": RELEASE,
            "language": "ko-KR",
            "assets": assets,
        })
    manifest = {
        "schema_version": 1,
        "release_version": RELEASE,
        "released_at": RELEASED_AT,
        "language": "ko-KR",
        "documents": documents,
    }
    payload = json.dumps(manifest, ensure_ascii=False, indent=2, sort_keys=False) + "\n"
    (ROOT / "docs" / "manual" / "release-manifest.json").write_text(payload, encoding="utf-8", newline="\n")
    (PUBLIC / "manifest.json").write_text(payload, encoding="utf-8", newline="\n")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("mode", choices=("docx", "manifest"))
    args = parser.parse_args()
    if args.mode == "docx":
        build_docx()
    else:
        build_manifest()


if __name__ == "__main__":
    main()
